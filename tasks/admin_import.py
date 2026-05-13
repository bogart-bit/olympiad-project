import re
import io
from django.shortcuts import render, redirect
from django.contrib import messages


# ─── ПАРСЕР ЗАДАЧ ────────────────────────────────────────────────────────────

def parse_tasks_docx(file_bytes):
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    tasks = []
    current_year = None
    current_grade = None
    current_task = None

    i = 0
    while i < len(lines):
        line = lines[i]

        if re.match(r'^\d+\s*класс', line, re.IGNORECASE):
            current_grade = line
            i += 1
            continue

        if re.fullmatch(r'\d{4}', line):
            current_year = int(line)
            i += 1
            continue

        if re.fullmatch(r'\d{1,2}', line):
            if current_task:
                tasks.append(current_task)
            current_task = {
                'year': current_year,
                'grade': current_grade or '',
                'condition': '',
                'correct_answer': '',
                'solution': '',
            }
            i += 1
            continue

        if current_task is None:
            i += 1
            continue

        if line.startswith('Тема:') or line.startswith('Нужные формулы:'):
            i += 1
            continue

        def _stop(l):
            return (
                any(l.startswith(k) for k in ['Условие:', 'Условия:', 'Решение:', 'Ответ:', 'Тема:', 'Нужные формулы:'])
                or bool(re.fullmatch(r'\d{4}', l.strip()))
                or bool(re.fullmatch(r'\d{1,2}', l.strip()))
                or bool(re.match(r'^\d+\s*класс', l, re.IGNORECASE))
            )

        if line.startswith('Условие:') or line.startswith('Условия:'):
            key_len = 9 if line.startswith('Условия:') else 8
            text = line[key_len:].strip()
            i += 1
            while i < len(lines) and not _stop(lines[i]):
                text += '\n' + lines[i]
                i += 1
            current_task['condition'] = text.strip()
            continue

        if line.startswith('Решение:'):
            text = line[8:].strip()
            i += 1
            while i < len(lines) and not _stop(lines[i]):
                text += '\n' + lines[i]
                i += 1
            current_task['solution'] = text.strip()
            continue

        if line.startswith('Ответ:'):
            text = line[6:].strip()
            i += 1
            while i < len(lines) and not _stop(lines[i]):
                text += '\n' + lines[i]
                i += 1
            current_task['correct_answer'] = text.strip()
            continue

        i += 1

    if current_task:
        tasks.append(current_task)

    return tasks


# ─── ПАРСЕР ФОРМУЛ ───────────────────────────────────────────────────────────

def parse_formulas_docx(file_bytes):
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    body = doc.element.body
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    current_year = None
    current_grade = None
    formulas = []

    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(t.text or '' for t in child.iter() if t.tag.endswith('}t')).strip()
            if re.match(r'^\d+\s*класс', text, re.IGNORECASE):
                current_grade = text
            elif re.fullmatch(r'\d{4}', text):
                current_year = int(text)
        elif tag == 'tbl':
            rows = child.findall(f'.//{{{NS}}}tr')
            order = 1
            for row_idx, row in enumerate(rows):
                if row_idx == 0:
                    continue
                cells = row.findall(f'{{{NS}}}tc')
                if len(cells) >= 2:
                    title = ''.join(t.text or '' for t in cells[0].iter() if t.tag.endswith('}t')).strip()
                    content = ''.join(t.text or '' for t in cells[1].iter() if t.tag.endswith('}t')).strip()
                    if title:
                        formulas.append({
                            'title': title,
                            'content': content,
                            'order': order,
                            'year': current_year,
                            'grade': current_grade or '',
                        })
                        order += 1

    return formulas


# ─── VIEW ────────────────────────────────────────────────────────────────────

def import_tasks_view(request):
    from tasks.models import Task
    return _import_view(request, 'tasks', Task, parse_tasks_docx,
                        'admin/import_tasks.html', 'admin:tasks_task_changelist')


def import_formulas_view(request):
    from tasks.models import Formula
    return _import_view(request, 'formulas', Formula, parse_formulas_docx,
                        'admin/import_formulas.html', 'admin:tasks_formula_changelist')


def _import_view(request, session_key, Model, parser_func, template, redirect_url):
    context = {'title': 'Импорт из файла', 'has_permission': True}

    if request.method == 'GET':
        return render(request, template, context)

    action = request.POST.get('action', 'preview')

    if action == 'preview':
        print("FILE:", request.FILES)
        uploaded_file = request.FILES.get('file')
        if not uploaded_file:
            messages.error(request, 'Файл не выбран.')
            return render(request, template, context)

        file_bytes = uploaded_file.read()
        if not uploaded_file.name.lower().endswith('.docx'):
            messages.error(request, 'Поддерживается только .docx файл.')
            return render(request, template, context)

        try:
            items = parser_func(file_bytes)
        except Exception as e:
            print("ОШИБКА ПАРСЕРА:", e)
            messages.error(request, f'Ошибка при чтении файла: {e}')
            return render(request, template, context)

        print("НАЙДЕНО:", len(items))
        if len(items) > 0:
            print("ПЕРВЫЙ:", items[0])

        if not items:
            messages.warning(request, 'Ничего не найдено. Проверь структуру файла.')
            return render(request, template, context)

        request.session[session_key] = items
        context['items'] = items
        context['count'] = len(items)
        return render(request, template, context)

    elif action == 'save':
        items = request.session.pop(session_key, [])
        if not items:
            messages.error(request, 'Нет данных. Загрузи файл снова.')
            return redirect(redirect_url)

        created = 0
        for item in items:
            Model.objects.create(**{k: v for k, v in item.items() if v is not None})
            created += 1

        messages.success(request, f'✅ Добавлено {created} записей!')
        return redirect(redirect_url)

    elif action == 'cancel':
        request.session.pop(session_key, None)
        return redirect(redirect_url)

    return render(request, template, context)
